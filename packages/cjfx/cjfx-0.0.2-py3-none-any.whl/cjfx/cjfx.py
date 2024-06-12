#!/bin/python3

'''
This module is a collection of functions and
classes which Celray James frequently uses with
Python 3.8+ it is ideal for quick prototyping and
data science projects.

Author  : Celray James CHAWANDA
Email   : celray.chawanda@outlook.com
Licence : MIT 2023
Repo    : https://github.com/celray

Updated : 26/12/2022
'''

# imports
import os
import sys
from collections import defaultdict
import ctypes
import base64
import datetime
import gzip
import math
import string
import multiprocessing
import pickle
import platform
import random
import shutil
import subprocess
import sys
import time
import warnings
import xml.etree.ElementTree as ET
import zipfile
from glob import glob
from itertools import product
from shutil import copyfile, move, rmtree
from pytube import YouTube, Playlist
import geopandas
import matplotlib.pyplot as plt
import numpy
import pandas
import requests
import xarray
from PIL import Image
from pySmartDL import SmartDL
from tqdm import tqdm
from osgeo import gdal, gdalconst, ogr
from docx import Document
from pydub import AudioSegment
from shapely.geometry import Point, Polygon
from docxcompose.composer import Composer

# functions
def points_to_geodataframe(point_pairs_list, columns = ['latitude', 'longitude'], auth = "EPSG", code = '4326', out_shape = '', format = 'gpkg', v = False, get_geometry_only = False):
    df = pandas.DataFrame(point_pairs_list, columns = columns)
    geometry = [Point(xy) for xy in zip(df['latitude'], df['longitude'])]

    if get_geometry_only:
        return geometry[0]

    gdf = geopandas.GeoDataFrame(point_pairs_list, columns = columns, geometry=geometry)
    drivers = {'gpkg': 'GPKG', 'shp': 'ESRI Shapefile'}

    gdf = gdf.set_crs(f'{auth}:{code}')
    
    if out_shape != '':
        if v: print(f'creating shapefile {out_shape}')
        gdf.to_file(out_shape, driver=drivers[format])
    
    return gdf


def convert_webm_to_mp3(webm_path, mp3_path):
    # load webm file
    print(f"converting {webm_path} to {mp3_path}")
    audio = AudioSegment.from_file(webm_path, format=webm_path.split(".")[-1])

    # save as mp3
    audio.export(mp3_path, format="mp3")



def download_video_youtube(url, path, audio_only = False, download_list = False, numbering = False):
    # create YouTube object

    is_playlist = False

    if "list=" in url:
        is_playlist = True
    
    video_urls = [url]
    if is_playlist:
        if download_list:
            playlist = Playlist(url)
            video_urls = playlist.video_urls

    counting = 0
    for vid_url in video_urls:
        counting += 1
            
        yt = YouTube(vid_url)
        
        # select the highest quality stream
        if audio_only:
            print(f"downloading audio from video: {vid_url}")
            stream = yt.streams.get_audio_only()
        else:
            print(f"downloading video: {vid_url}")
            stream = yt.streams.get_highest_resolution()
        
        create_path(f"{path}/")

        filename = stream.download(path, f"{counting}. {yt.title}.mp4" if numbering else f"{yt.title}.mp4")
        full_path = os.path.join(path, filename)

        if audio_only:
            convert_webm_to_mp3(full_path, full_path.replace(full_path.split(".")[-1], "mp3"))
            delete_file(full_path, v=False)

def get_relative_path(base_path, target_path):
    return os.path.relpath(target_path, base_path)


def install_package(package_name, version = None):
    if platform.system() == "Windows":
        try: os.system(f"pip3 install {package_name}{'==' + version if not version is None else ''}")
        except: raise

    else:
        install_command = f"pip3 install {package_name}{'==' + version if not version is None else ''} --break-system-packages"
        print(install_command)
        os.system(install_command)
        

def get_usgs_timeseries(file_path, monthly=True, convert=True):
    df = pandas.read_csv(file_path, delim_whitespace=True, names=['date', 'time', 'flow'])
    df['date'] = pandas.to_datetime(df['date'], format='%Y%m%d')
    
    if convert:
        df['flow'] = df['flow'] * 0.0283168

    if monthly:
        df = df.groupby(pandas.Grouper(key='date', freq='M')).mean()
        df.reset_index(inplace=True)
        df.rename(columns={'index': 'date'}, inplace=True)

    return df[['date', 'flow']]


def decode_64(encoded:str) -> str:
    return base64.b64decode(encoded).decode('utf-8')

def wait(duration:int = 1):
    time.sleep(duration)

def ignore_warnings(ignore:bool = True, v:bool = False):
    if ignore:
        warnings.filterwarnings("ignore")
        if v: print("warnings ignored")
    else:
        warnings.filterwarnings("default")
        if v: print("warnings not ignored")

def get_raster_value_for_coords(lat, lon, rasterio_array):
    row, col = rasterio_array.index(lon,lat)
    val = rasterio_array.read(1)[row,col]
    return(val)

def rasterise_shape(shape_file:str, column_name:str, destination_tif:str, template_tif:str, no_data:int = -999) -> None:
    data = gdal.Open(template_tif, gdalconst.GA_ReadOnly)
    geo_transform = data.GetGeoTransform()
    x_min = geo_transform[0]
    y_max = geo_transform[3]
    x_max = x_min + geo_transform[1] * data.RasterXSize
    y_min = y_max + geo_transform[5] * data.RasterYSize
    x_res = data.RasterXSize
    y_res = data.RasterYSize
    mb_v = ogr.Open(shape_file)
    mb_l = mb_v.GetLayer()
    pixel_width = geo_transform[1]
    target_ds = gdal.GetDriverByName('GTiff').Create(destination_tif, x_res, y_res, 1, gdal.GDT_Int32)
    target_ds.SetGeoTransform(data.GetGeoTransform())
    target_ds.SetProjection(data.GetProjection())
    # target_ds.SetGeoTransform((x_min, pixel_width, 0, y_min, 0, pixel_width))
    band = target_ds.GetRasterBand(1)
    band.SetNoDataValue(no_data)
    band.FlushCache()
    gdal.RasterizeLayer(target_ds, [1], mb_l, options=[f"ATTRIBUTE={column_name}"])

    target_ds = None

def extract_timeseries_from_netcdf(fn, variable, lon_coord, lat_coord, method = 'nearest', v = False):
    if v: print(f'getting a {lon_coord},{lat_coord} timeseries for {variable}')
    ds = xarray.open_dataset(fn)
    dsloc = ds.sel(lon=lon_coord, lat=lat_coord, method=method)
    return dsloc[variable].to_dataframe()

def rand_apha_num(length = 8):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

def make_plot(data_pd, x_col:str, y1_cols:list, y1_axis_label, save_path:str, title = None, y2_axis_label = None, y2_cols:list = None,
                x_axis_label:str=None, y1_labels:list = None, y2_labels:list = None, x_limits:tuple = None, y1_limits:tuple = None, y2_limits:tuple = None,
                show_plot = False, legend = False, secondary_type = 'bar', chart_size:tuple = (14, 5)):
    '''
    x_limits : example (0, 10)
    y_labels : example ['ET', 'PCP'] # no need to be same as y_cols but y_cols are used if this is not given

    secondary_type: 'bar', 'line'
    '''

    import seaborn
    from matplotlib import pyplot as plt

    plt.close()
    plt.margins(x=None, y=None, tight=True)
    plt.tight_layout(pad = 0)

    ax1 = seaborn.set_style(style=None, rc=None )

    fig = plt.figure(figsize = chart_size, dpi=300)

    # make axis
    ax1  = fig.add_subplot(1, 1, 1)

    min_ax1_ini =  1000000000000
    max_ax1_ini = -1000000000000
    for index in range(0, len(y1_cols)):
        ax1.plot(data_pd[x_col].tolist(), data_pd[y1_cols[index]].tolist(), label = y1_labels[index] if not y1_labels is None else y1_cols[index])

        if min(data_pd[y1_cols[index]].tolist()) < min_ax1_ini: min_ax1_ini = min(data_pd[y1_cols[index]].tolist())
        if max(data_pd[y1_cols[index]].tolist()) > max_ax1_ini: max_ax1_ini = max(data_pd[y1_cols[index]].tolist())

    min_ax1 = max(0, min_ax1_ini - (max_ax1_ini - min_ax1_ini)/5 )
    max_ax1 = max_ax1_ini + (max_ax1_ini - min_ax1_ini)/1.2

    min_ax1 = round(min_ax1/10)*10 if min_ax1 < 80 else (round(min_ax1/100)*100 if min_ax1 < 800 else round(min_ax1/1000)*1000)

    if min_ax1 > min_ax1_ini:
        difference_ax1 = round(1 + (min_ax1 - min_ax1_ini)/10) * 10 if (min_ax1 - min_ax1_ini) < 80 else (round(1 + (min_ax1 - min_ax1_ini)/100)*100 if (min_ax1 - min_ax1_ini) < 800 else round(1 + (min_ax1 - min_ax1_ini)/1000)*1000)
        min_ax1 = min_ax1 - difference_ax1

    difference_ax1 = round(1 + (max_ax1 - min_ax1)//10) * 10 if (max_ax1 - min_ax1) < 80 else (round(1 + (max_ax1 - min_ax1)//100)*100 if (max_ax1 - min_ax1) < 800 else round(1 + (max_ax1 - min_ax1)/1000)*1000)

    max_ax1 = min_ax1 + difference_ax1

    # set axes labels
    ax1.set_ylabel(y1_axis_label)
    if not title is None: ax1.set_title(title, fontdict={'fontsize': 18})
    if not x_axis_label is None: ax1.set_xlabel(x_axis_label)

    if y1_limits is None:
        ax1.set_ylim([min_ax1, max_ax1])
        ax1.grid(axis='y', which='major', color = 'lightgrey')
        ax1.set_yticks([min_ax1 + (i * (max_ax1 - min_ax1) / 5) for i in range(1, 6)])
    else: ax1.set_ylim([y1_limits[0], y1_limits[1]])

    if not x_limits is None: ax1.set_xlim([x_limits[0], x_limits[1]])


    if not y2_cols is None:
        ax2 = ax1.twinx()

        min_ax2_ini =  1000000000000
        max_ax2_ini = -1000000000000

        if secondary_type == 'bar':
            for index in range(0, len(y2_cols)):
                ax2.bar(data_pd[x_col].tolist(), data_pd[y2_cols[index]].tolist(), 0.6, label = y2_labels[index] if not y2_labels is None else y2_cols[index], color = 'black')
                
                if min(data_pd[y2_cols[index]].tolist()) < min_ax2_ini: min_ax2_ini = min(data_pd[y2_cols[index]].tolist())
                if max(data_pd[y2_cols[index]].tolist()) > max_ax2_ini: max_ax2_ini = max(data_pd[y2_cols[index]].tolist())
        
        elif secondary_type == 'line':
            for index in range(0, len(y2_cols) + 1):
                ax2.plot(data_pd[x_col].tolist(), data_pd[y2_cols[index]].tolist(), label = y2_labels[index] if not y2_labels is None else y2_cols[index])
                
                if min(data_pd[y2_cols[index]].tolist()) < min_ax2_ini: min_ax2_ini = min(data_pd[y2_cols[index]].tolist())
                if max(data_pd[y2_cols[index]].tolist()) > max_ax2_ini: max_ax2_ini = max(data_pd[y2_cols[index]].tolist())

        else: raise ValueError('\'secondary_type\' can only be either \'bar\' or \'line\'')

        # label axes
        if not y2_axis_label is None: ax2.set_ylabel(y2_axis_label)

        min_ax2 = max(0, min_ax2_ini - (max_ax2_ini - min_ax2_ini)/5)
        max_ax2 = max_ax2_ini + (max_ax2_ini - min_ax2_ini)/0.9

        min_ax2 = round(min_ax2/10)*10 if min_ax2 < 80 else (round(min_ax2/100)*100 if min_ax2 < 800 else round(min_ax2/1000)*1000)

            
        if min_ax2 > min_ax2_ini:
            difference_ax2 = round(1 + (min_ax2 - min_ax2_ini)/10) * 10 if (min_ax2 - min_ax2_ini) < 80 else (round(1 + (min_ax2 - min_ax2_ini)/100)*100 if (min_ax2 - min_ax2_ini) < 800 else round(1 + (min_ax2 - min_ax2_ini)/1000)*1000)
            min_ax2 = min_ax2 - difference_ax2

        difference_ax2 = round(1 + (max_ax2 - min_ax2)//10) * 10 if (max_ax2 - min_ax2) < 80 else (round(1 + (max_ax2 - min_ax2)//100)*100 if (max_ax2 - min_ax2) < 800 else round(1 + (max_ax2 - min_ax2)/1000)*1000)

        max_ax2 = min_ax2 + difference_ax2




        # invert axis (need to set limits too)
        if y2_limits is None:
            ax2.set_yticks([i * (max_ax2 - min_ax2) / 5 for i in range(1, 6)])
            ax2.set_ylim([min_ax2, max_ax2][::-1])
        else:
            ax2.set_ylim([y2_limits[0], y2_limits[1]][::-1])

    handles,labels = [],[]
    for ax in fig.axes:
        for h,l in zip(*ax.get_legend_handles_labels()):
            handles.append(h)
            labels.append(l)

    if legend: plt.legend(handles, labels, loc="center right",)

    plt.savefig(save_path)
    if show_plot: plt.show()

    return plt

def goto_dir(obj_):
    '''
    obj_ should be __file__
    '''
    me = os.path.realpath(obj_)
    os.chdir(os.path.dirname(me))

def smart_copy(source_dir, destination_dir, move = False):
    destination_dir = destination_dir if not destination_dir.endswith("/") else destination_dir[:-1]
    source_dir = source_dir if not source_dir.endswith("/") else source_dir[:-1]
    splitter = os.path.basename(source_dir)

    all_files_names = list_all_files(source_dir)
    class file_paths:
        def __init__(self, src, dst):
            self.src = src
            self.dst = dst

    sort_size = []
    files_dict = {}

    print("\n\tsorting files")
    for fn in all_files_names:
        size = os.path.getsize(fn)
        while size in sort_size:
            size += random.random()
        sort_size.append(size)
        files_dict[size] = file_paths(src = fn, dst = "{0}/{1}{2}".format(destination_dir, splitter, fn.split(splitter)[-1]))

    sort_size.sort(reverse=True)

    count = 0
    end = len(sort_size)
    for sorted_file in sort_size:
        count += 1
        show_progress(count = count, end_val=end, string_after= "{2} {0} - {1} MB".format(
            os.path.basename(files_dict[sorted_file].dst).lower(),
            round(sorted_file/1000000, 3),
            "copying" if not move else "moving"
            ))

        if not os.path.isfile(files_dict[sorted_file].src):
            print("\nfile not found")
            return

        copy_file(files_dict[sorted_file].src, files_dict[sorted_file].dst, delete_source = move)

    print("")

def print_dict(dictionary, columns = 4, width_k = 10, width_val = 12, sep = ":") -> None:
    counter = columns

    keys = [k for k in dictionary]
    keys.sort()

    print_str = ""
    for key in keys:
        print_str += f"{str(key).rjust(width_k)} {sep} {str(dictionary[key]).ljust(width_val)}     "
        counter -= 1

        if counter == 0:
            counter = columns
            print(print_str)
            print_str = ""
        
    print(print_str)
    return print_str


def transparent_image(in_file, out_file, threshold = 251, v = True):
    img = Image.open(in_file)
    img = img.convert("RGBA")

    pixdata = img.load()

    width, height = img.size
    for y in range(height):
        for x in range(width):
            if pixdata[x, y][0] > threshold:
                if pixdata[x, y][1]  > threshold:
                    if pixdata[x, y][2]  > threshold:
                        pixdata[x, y] = (255, 255, 255, 0)

    img.save(out_file, "PNG")
    if v:
        print(f"\t\t - saved transparent image to {out_file}")

def get_file_size(file_path):
    return float(os.path.getsize(file_path))/1012

def cd(path):
    if os.path.isdir(path):
        os.chdir(path)
    else:
        sys.exit("\t! the path {0} does not exist".format(path))

def slope_intercept(x1,y1,x2,y2):
    a = (y2 - y1) / (x2 - x1)
    b = y1 - a * x1     
    return a,b

def distance(coords_a, coords_b):
    return math.sqrt(((coords_b[0] - coords_a[0]) ** 2) + (coords_b[1] - coords_a[1]) ** 2)

def merge_documents(out_fn, list_of_in_fn, from_empty = False, v = False):
    if len(list_of_in_fn) < 2:
        print('\t! you need atleast two documents')
        return

    create_path(out_fn)

    if from_empty:
        # initialise empty document
        document = Document(); document.save(out_fn)
    else:
        if not list_of_in_fn[0] == out_fn:
            copy_file(list_of_in_fn[0], out_fn)
        else:
            if v: print(f"\t ! appending to the existing out file")

        list_of_in_fn = list_of_in_fn[1:]

    master = Document(out_fn)
    composer = Composer(master)

    for doc_fn in list_of_in_fn:
        composer.append(Document(doc_fn))

    master.save(out_fn)  

def resize_image(in_file:str, out_file:str, size:int = None, ratio = None, width = None, height = "auto"):
    fsize = get_file_size(in_file)

    if not size is None:
        ratio = 1

        class point:
            def __init__(self, X, Y) -> None:
                self.X = X
                self.Y = Y

        history = []
        # print(abs(fsize-size))
        while fsize > size:
            # if len(history) > 1:
            #     a, b = slope_intercept(history[-2].X, history[-2].Y, history[-1].X, history[-1].Y)
            #     ratio = b/a
            # else:
            #     ratio -= 0.2
            # if ratio < 0.05: ratio = random.randint(1, 9) * random.randint(1, 9) * 0.01
            image = Image.open(in_file)

            old_width, old_height = image.size
            newsize = (int(old_width * ratio), int(old_height * ratio))

            image = image.resize(newsize)
            create_path(out_file)
            image.save(out_file)

            fsize = get_file_size(out_file)
            ratio -= 0.01

            sys.stdout.write('\r\t> testing ratio: ' + str(ratio))
            sys.stdout.flush()

            history.append(point(ratio, fsize-size))

        print(f'\n\t> reduced by {ratio}\n')
        
    else:
        image = Image.open(in_file)

        old_width, old_height = image.size

        if ratio is None:
            if width is None:
                print("You should supply either width or ratio")
                return

            if height == "auto":
                newsize = (width, int((old_height/old_width) * width))
            else:
                newsize = (width, height)
            
        else:
            newsize = (int(old_width * ratio), int(old_height * ratio))

        image = image.resize(newsize)
        create_path(out_file)
        image.save(out_file)
        return True
    
def download_file(url, dest_dir, replace = True, v = True):

    create_path(dest_dir)
    if v: print(f"downloading {url}...")
    dest_dir = f"{dest_dir}/" if not dest_dir.endswith("/") else dest_dir
    if not replace:
        if exists(f"{dest_dir}{url.split('/')[-1]}".replace('%20', ' ')):
            return None

    obj = SmartDL(url, dest_dir)
    try:
        obj.start()
        path = obj.get_dest()
        return path
    except:
        return None


def download_file2(url, save_path, overwrite = False, v=True):
    if v:
        print(f"\ndownloading {url}")
    fname = file_name(url, extension=True)
    save_dir = os.path.dirname(save_path)
    save_fname = "{0}/{1}".format(save_dir, fname)

    response = requests.get(url, stream=True)
    if not os.path.isdir(save_dir):
        os.makedirs(save_dir)
    
    if overwrite:
        if exists(save_fname):
            delete_file(save_fname)
    else:
        if exists(save_fname):
            error("cannot download file; it already exists on disk")
            return

    with open(save_fname, "wb") as handle:
        for data in tqdm(response.iter_content()):
            handle.write(data)


def apply_parallel(function_name, number_of_processes, *args):
    "has issues"
    with multiprocessing.Pool(processes=number_of_processes) as pool:
        results = pool.starmap(function_name, product(args[0], args[1]))

    return results


def xml_children_attributes(xml_file_name, x_path):
    root = ET.parse(xml_file_name).getroot()
    result = {}
    for element in root.findall(x_path):
        for child in element:
            result[child.tag] = child.text

    return result


def python_variable(option, filename, variable=None):
    '''
    option: save, load or open

    '''
    if (option == "save") and (variable is None):
        print("\t! please specify a variable")

    if option == "save":
        create_path(filename)
        with open(filename, 'wb') as f:
            pickle.dump(variable, f)

    if (option == "load") or (option == "open"):
        with open(filename, "rb") as f:
            variable = pickle.load(f)

    return variable

def get_nse(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    """
    This function calculates the NSE from columns in a pandas dataframe. The dataframe
    should have two columns, one for the simulated data and one for the observed data.
    The columns should be named as the arguments of this function. The dataframe may not
    have a date but it is recommended to have a date column.
    """
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    try:
        nse = 1 - sum((df_clean[observed_column] - df_clean[simulated_column]) ** 2) / \
            sum((df_clean[observed_column] - numpy.mean(df_clean[observed_column])) ** 2)
    except ZeroDivisionError:
        nse = None

    return nse

def get_pbias(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    """
    This function calculates the PBIAS from columns in a pandas dataframe. The dataframe
    should have two columns, one for the simulated data and one for the observed data.
    The columns should be named as the arguments of this function. The dataframe may not
    have a date but it is recommended to have a date column.
    """
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    try:
        pbias = sum((df_clean[observed_column] - df_clean[simulated_column])) / \
            sum(df_clean[observed_column]) * 100
    except ZeroDivisionError:
        pbias = None
    return pbias


def report(string, printing=False):
    if printing:
        print(f"\t> {string}")
    else:
        sys.stdout.write("\r" + string)
        sys.stdout.flush()


def print_list(list_object):
    print("\t> list content")
    for item in list_object:
        print("\t  - {0}".format(item))


def disp(string_):
    print(f"\t- {string_}")


def create_path(path_name, v = False):
    path_name = os.path.dirname(path_name)
    if path_name == '':
        path_name = './'
    if not os.path.isdir(path_name):
        os.makedirs(path_name)
        if v:
            print(f"\t> created path: {path_name}")
    
    return path_name

def flow_duration_curve(df, col_name):
    # Select the data for the specified type
    flow_data = df[col_name]

    # Sort the flow data in decreasing order
    sorted_data = flow_data.sort_values(ascending=False)
    
    # Calculate the rank of each data point
    ranks = numpy.arange(1,len(sorted_data)+1)
    
    # Calculate the exceedance probability
    exceedance_prob = (ranks/(len(sorted_data)+1))*100
    
    # Create a DataFrame for plotting
    df_plot = pandas.DataFrame({
        'Exceedance Probability (%)': exceedance_prob,
        col_name: sorted_data
    }).reset_index(drop=True)
    
    return df_plot


def delete_path(path, v=True):
    try:
        if v:
            rmtree(path)
            print(f"\t! {path} has been deleted")
    except:
        print("\t! could not delete the diractory at {path}".format(path=path))


def is_file(path_):
    if os.path.isfile(path_):
        return True
    else:
        return False


def resample_ts_df(df, column_name, t_step="M", resample_type="mean", only_numeric = True):
    '''
    resample_type   : mean, sum
    t_step          : M - > month
                      Y - > year
                      W - > week
                      3T- > 3 minutes
                      3S- > 3 seconds
    '''
    df[column_name] = pandas.to_datetime(df[column_name])
    if resample_type == "mean":
        result_df = df.resample(t_step, on = column_name).mean(numeric_only = only_numeric)
    elif resample_type == "sum":
        result_df = df.resample(t_step, on = column_name).sum(numeric_only = only_numeric)
    else:
        print("\t ! please select a resample type: available > mean, sum")
        sys.exit()
    return result_df


def raster_statistics(tif_file, v =True):
    ds = gdal.Open(tif_file)
    minimum, maximum, mean, std_dev = ds.GetRasterBand(1).GetStatistics(0, 1)

    class gdal_stats:
        def __init__(self, mn, mx, mean, std_dev):
            self.minimum = mn
            self.maximum = mx
            self.mean = mean
            self.stdev = std_dev

        def __repr__(self):
            return 'min: {0}, max: {1}, mean: {2}, sdev: {3}'.format(self.minimum, self.maximum,  self.mean, self.stdev)
    if v:
        report(f"processing {file_name(tif_file, extension=True)}")

    all_stats = gdal_stats(minimum, maximum, mean, std_dev)
    return all_stats


def empty_line():
    print("")

def hide_folder(path_to_folder):
    ctypes.windll.kernel32.SetFileAttributesW(path_to_folder, 2)

def time_stamp():
    return datetime.datetime.now()


def copy_folder(src_dir, dst_dir, delete_source = False, v = False, progress = True, replace = True, exception_list = None):
    all_files = list_all_files(src_dir)

    current = 0; finish = len(all_files)
    for fn in all_files:
        current += 1
        
        skip_file = False
        if not exception_list is None:
            for exception_item in exception_list:
                if exception_item in fn:
                    skip_file = True

        new_fn = f"{dst_dir}/{get_relative_path(src_dir, fn)}"

        if progress and v:
            if skip_file:
                show_progress(current, finish, scroll_text=f"{'skipping'} {fn}")
            else:
                show_progress(current, finish, scroll_text=f"{'copying' if not delete_source else 'moving'} {fn}")
                copy_file(fn, new_fn, replace=replace, v = False, delete_source=delete_source)
        elif progress:
            if not skip_file:
                show_progress(current, finish)
                copy_file(fn, new_fn, replace=replace, v = False, delete_source=delete_source)

    if delete_source:
        delete_path(src_dir, v=False)


def delete_file(file_name, v = True):
    if os.path.isfile(file_name):
        try:
            os.remove(file_name)
            if v: print("\t> deleted file {fn}".format(fn=file_name))
        except:
            print("\t! could not delete file {fn}".format(fn=file_name))
    else:
        if v: print("\t! the file, {fn}, does not exist".format(fn=file_name))


def run_swat_plus(txtinout_dir, final_dir, executable_path, v = True):
    os.chdir(txtinout_dir)
    
    if not v:
        # Run the SWAT+ but ignore output and errors
        subprocess.run([executable_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    else:
        
        yrs_line = read_from('time.sim')[2].strip().split()

        yr_from = int(yrs_line[1])
        yr_to = int(yrs_line[3])

        delta = datetime.datetime(yr_to, 12, 31) - datetime.datetime(yr_from, 1, 1)

        CREATE_NO_WINDOW = 0x08000000

        if platform.system() == "Windows":
            process = subprocess.Popen(executable_path, stdout=subprocess.PIPE, creationflags=CREATE_NO_WINDOW )
        else:
            process = subprocess.Popen(executable_path, stdout=subprocess.PIPE)

        counter = 0

        current = 0
        number_of_days = delta.days + 1

        day_cycle = []
        previous_time = None

        while True:
            line = process.stdout.readline()
            line_parts = str(line).strip().split()
            if not "Simulation" in line_parts: pass
            elif 'Simulation' in line_parts:
                ref_index = str(line).strip().split().index("Simulation")
                year = line_parts[ref_index + 3]
                month = line_parts[ref_index + 1]
                day = line_parts[ref_index + 2]


                month = f"0{month}" if int(month) < 10 else month
                day = f"0{day}" if int(day) < 10 else day
                
                current += 1
                
                if not previous_time is None:
                    day_cycle.append(datetime.datetime.now() - previous_time)

                if len(day_cycle) > 40:
                    if len(day_cycle) > (7 * 365.25):
                        del day_cycle[0]

                    av_cycle_time = sum(day_cycle, datetime.timedelta()) / len(day_cycle)
                    eta = av_cycle_time * (number_of_days - current)

                    eta_str = f"  ETA - {format_timedelta(eta)}:"
                    

                else:
                    eta_str = ''

                show_progress(current, number_of_days, bar_length=20, string_after= f'  >> current date: {day}/{month}/{year} - final-date: 31/12/{yr_to} {eta_str}')

                previous_time = datetime.datetime.now()
            elif "ntdll.dll" in line_parts:
                print("\n! there was an error running SWAT+\n")
            if counter < 10:
                counter += 1
                continue

            if len(line_parts) < 2: break

        show_progress(1, 1, string_after= f'                                                                                      ')
        print("\n")
    
    os.chdir(final_dir)


def open_file_in_code(file_path):
    '''
    This function opens a file in the code editor
    '''
    if file_path == "":
        file_path = os.path.join(os.getcwd(), __file__)
    os.system(f"code {file_path}")



def format_timedelta(delta: datetime.timedelta) -> str:
    """Formats a timedelta duration to [N days] %H:%M:%S format"""
    seconds = int(delta.total_seconds())

    secs_in_a_day = 86400
    secs_in_a_hour = 3600
    secs_in_a_min = 60

    days, seconds = divmod(seconds, secs_in_a_day)
    hours, seconds = divmod(seconds, secs_in_a_hour)
    minutes, seconds = divmod(seconds, secs_in_a_min)

    time_fmt = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

    if days > 0:
        suffix = "s" if days > 1 else ""
        return f"{days} day{suffix} {time_fmt}"
    else:
        return f"{time_fmt}"



def show_progress(progress, end, dt=None, string_before="", string_after="", bar_length=100, precision = 1, d_count = None, scroll_text = None):
    '''
    dt: timedelta from datetime
    '''

    if platform.system() != "Windows":
        if (os.get_terminal_size(0)[0] - (len(string_after) + 21)) < bar_length:
            bar_length = os.get_terminal_size(0)[0] - (len(string_after) + 21)

    if bar_length < 5: bar_length = 5

    if not scroll_text is None:

        if platform.system() != "Windows":
            sys.stdout.write('\r' + str(scroll_text).ljust(os.get_terminal_size(0)[0]))
        else:
            sys.stdout.write('\r' + str(scroll_text).ljust(max(151, 0)))

        sys.stdout.flush(); print()

    percent = float(progress) / end
    hashes = "█" * int(round(percent * bar_length))
    spaces = '░' * (bar_length - len(hashes))
    eta = 0
    if dt is not None:

        if len(dt) > 5:

            cycle_time = sum(dt, datetime.timedelta()) / (len(dt) if d_count is None else d_count)
            cycles_to_go = end - progress
            eta = format_timedelta((cycle_time * cycles_to_go))

        else:
            dt = None

    sys.stdout.write("\r{str_b}{bar} {sp}{pct}% {str_after}          ".format(
        str_b=string_before, sp = "  " if percent < 100 else "",
        bar=hashes + spaces,
        pct='{:06.2f}'.format(percent * 100),
        str_after=string_after if dt is None else string_after + " - " + "eta: " + eta))
    sys.stdout.flush()




def confirm():
    print("\n\t>> set line reached, terminating program")
    answer = input(
        "\t>> press ENTER to continue, 'N' and ENTER to terminate: ")
    if answer == "N":
        sys.exit()


def open_tif_as_array(tif_file, big_tif = True, band = 1):
    if big_tif:
        dataset = gdal.Open(tif_file, gdal.GA_ReadOnly)
        band = dataset.GetRasterBand(band)
        return band.ReadAsArray()
    else:
        im = Image.open(tif_file)
        imarray = numpy.array(im)
        return imarray


def save_array_as_image(np_array, out_file, v = False):
    if not os.path.dirname(out_file) == "":
        create_path(os.path.dirname(out_file))
    im = Image.fromarray(np_array)
    if v:
        print(f'\t> writing array to {out_file}')
    im.save(out_file)

def create_icon(source_image_path, destination_path):
    # create_path(os.path.dirname(destination_path))
    img = Image.open(source_image_path)
    icon_sizes = [(16, 16), (32, 32), (48, 48), (64, 64)]
    img.save(destination_path if destination_path.endswith(".ico")
             else f"{destination_path}.ico", sizes=icon_sizes)

def unzip_file(f_name, destination_of_contents):
    show("extracting {0} to {1}".format(f_name, destination_of_contents))
    
    create_path(f'{destination_of_contents}/{file_name(f_name, extension=True)[:-3]}')
    if f_name.lower().endswith('.gz'):
        with gzip.open(f_name, 'rb') as f_in:
            with open(f'{destination_of_contents}/{file_name(f_name, extension=True)[:-3]}', 'wb') as f_out:
                shutil.copyfileobj(f_in, f_out)
    else:
        try:
            with zipfile.ZipFile(f_name, "r") as zip_ref:
                zip_ref.extractall(destination_of_contents)
        except:
            print(f"{f_name} is probably a bad zip file")


def single_spaces(string_):
    for i in range(0, 20):
        string_ = string_.replace("  ", " ")
    return string_

def view(file_name, ask = False):
    import platform
    
    if platform.system() == "Windows":
        if not ask:
            os.startfile(file_name)
        else:
            answer = input(f"\t Open {file_name} ?\n\t Y/N/E: ")
            if answer == "Y":
                os.startfile(file_name)
            if answer == "N":pass
            if answer == "E":
                sys.exit()

def quit():
    sys.exit("! mannually exited...")

def insert_newlines(input_str, word_count=11):
    words = input_str.split(' ')
    lines = [' '.join(words[i:i+word_count]) for i in range(0, len(words), word_count)]
    return '\n'.join(lines)


def write_to(filename, text_to_write, v=False, mode = "overwrite"):
    '''
    a function to write to file
    modes: overwrite/o; append/a
    '''
    try:
        if not os.path.isdir(os.path.dirname(filename)):
            os.makedirs(os.path.dirname(filename))
            if v:
                print("! the directory {0} has been created".format(
                    os.path.dirname(filename)))
    except:
        pass

    if (mode == "overwrite") or (mode == "o"):
        g = open(filename, 'w', encoding="utf-8")
    elif (mode == "append") or (mode == "a"):
        g = open(filename, 'a', encoding="utf-8")
    try:
        g.write(text_to_write)
        if v:
            print('\n\t> file saved to ' + filename)
    except PermissionError:
        print("\t> error writing to {0}, make sure the file is not open in another program".format(
            filename))
        response = input("\t> continue with the error? (Y/N): ")
        if response == "N" or response == "n":
            sys.exit()
    g.close


def exists(path_):
    if os.path.isdir(path_):
        return True
    if os.path.isfile(path_):
        return True
    return False

def gdal_datatypes():
    return {
            'Byte': gdal.GDT_Byte,
            'Int16': gdal.GDT_Int16,
            'Int32': gdal.GDT_Int32,
            'UInt16': gdal.GDT_UInt16,
            'UInt32': gdal.GDT_UInt32,
            'CInt16': gdal.GDT_CInt16,
            'CInt32': gdal.GDT_CInt32,
            'Float32': gdal.GDT_Float32,
            'Float64': gdal.GDT_Float64,
            'CFloat32': gdal.GDT_CFloat32,
            'CFloat64': gdal.GDT_CFloat64,
        }

def resample_raster(original_file:str, destination_file:str, resolution:float, authority = "ESRI", auth_code = '54003', resampleAlg = "Bilinear", data_type = "Int16", srcNodata = -32768, dstNodata = -999) -> bool:
    dtt = gdal_datatypes()

    report(f"\rresampling {original_file}                                             ")

    ds = gdal.Warp(destination_file, original_file, dstSRS=f'{authority}:{auth_code}', resampleAlg=f"{resampleAlg}", srcNodata=srcNodata, dstNodata=dstNodata, outputType=dtt[data_type], xRes=resolution, yRes=resolution)
    ds = None

    return True

def create_polygon_geodataframe(lat_list, lon_list, auth = 'EPSG', code = 4326):
    geometry_ = Polygon(zip(lon_list, lat_list))
    polygon = geopandas.GeoDataFrame(index=[0], crs=f"{auth}:{code}", geometry=[geometry_])

    return polygon

    
def reproject_raster(input_raster, epsg, output_raster, method = "mode"):
    command = "gdalwarp -overwrite -t_srs EPSG:" + str(epsg) + " -r " + method + " -of GTiff " + input_raster + " " + output_raster
    os.system(command)
    print("\t> reprojected {0} to {1}".format(input_raster, epsg))

def set_nodata(input_tif, output_tif, nodata = -999):
    command = f"gdal_translate -of GTiff -a_nodata {nodata} {input_tif} {output_tif}"
    os.system(command)
    print(f"\t> set no data from {input_tif} to {output_tif}")

def get_extents(raster):
    src = gdal.Open(raster)
    upper_lef_x, xres, xskew, upper_left_y, yskew, yres  = src.GetGeoTransform()
    lower_right_x = upper_lef_x + (src.RasterXSize * xres)
    lower_right_y = upper_left_y + (src.RasterYSize * yres)
    return upper_lef_x, lower_right_y, lower_right_x, upper_left_y

def file_name(path_, extension=True):
    if extension:
        fn = os.path.basename(path_)
    else:
        fn = os.path.basename(path_).split(".")[0]
    return(fn)

def read_from(filename, decode_codec = None, v=False):
    '''
    a function to read ascii files
    '''
    try:
        if not decode_codec is None: g = open(filename, 'rb')
        else: g = open(filename, 'r')
    except:
        print(
            "\t! error reading {0}, make sure the file exists".format(filename))
        return

    file_text = g.readlines()
    
    if not decode_codec is None: file_text = [line.decode(decode_codec) for line in file_text]

    if v:
        print("\t> read {0}".format(file_name(filename)))
    g.close
    return file_text


def error(text_):
    print("\t! {string_}".format(string_=text_))

def strip_characters(input_:str, characters:str):
    return ''.join([i for i in input_ if not i in characters]) 

def list_files(folder, extension="*"):
    if folder.endswith("/"):
        if extension == "*":
            list_of_files = glob(folder + "*")
        else:
            list_of_files = glob(folder + "*." + extension)
    else:
        if extension == "*":
            list_of_files = glob(folder + "/*")
        else:
            list_of_files = glob(folder + "/*." + extension)
    return list_of_files


def list_all_files(folder, extension="*"):
    list_of_files = []
    # Getting the current work directory (cwd)
    thisdir = folder

    # r=root, d=directories, f = files
    for r, d, f in os.walk(thisdir):
        for file in f:
            if extension == "*":
                list_of_files.append(os.path.join(r, file))
            elif "." in extension:
                if file.endswith(extension[1:]):
                    list_of_files.append(os.path.join(r, file))
                    # print(os.path.join(r, file))
            else:
                if file.endswith(extension):
                    list_of_files.append(os.path.join(r, file))
                    # print(os.path.join(r, file))

    return list_of_files


def list_folders(directory):
    """
    directory: string or pathlike object
    """
    all_dirs = os.listdir(directory)
    dirs = [dir_ for dir_ in all_dirs if os.path.isdir(
        os.path.join(directory, dir_))]
    return dirs


def plot(x_list, y_list, plot_type='line', x_label="", y_label="", title_=""):
    '''
    plot_type: line
    '''
    plt.plot(x_list, y_list)
    plt.title(title_)
    plt.xlabel(x_label)
    plt.ylabel(y_label)
    plt.show()


def show(text_, error=False, same_line=True):
    if error:
        if same_line:
            sys.stdout.write("\r\t! {string_}                              ".format(string_=text_))
            sys.stdout.flush()
        else:
            sys.stdout.write("\r\t! {string_}".format(string_=text_))
            sys.stdout.flush()
    else:
        if same_line:
            sys.stdout.write("\r\t> {string_}                               ".format(string_=text_))
            sys.stdout.flush()
        else:
            sys.stdout.write("\r\t> {string_}".format(string_=text_))
            sys.stdout.flush()



def copy_file(filename, destination_path, delete_source=False, v = False, replace = True):
    '''
    a function to copy files
    '''
    if not replace:
        if exists(destination_path):
            if v:
                print(f"\t - file exists, skipping")
            return

    if not exists(filename):
        if not v:
            return
        print("\t> The file you want to copy does not exist")
        print(f"\t    - {filename}\n")
        ans = input("\t> Press  E then ENTER to Exit or C then ENTER to continue: ")

        counter = 0
        while (not ans.lower() == "c") and (not ans.lower() == "e"):
            ans = input("\t> Please, press E then ENTER to Exit or C then ENTER to continue: ")
            if counter > 2:
                print("\t! Learn to read instrunctions!!!!")
            counter += 1

        if ans.lower() == 'e': quit()
        if ans.lower() == 'c':
            write_to("log.txt", f"{filename}\n", mode='append')
            return


    if v:
        if delete_source:
            print(f"\t - [{get_file_size(filename)}] moving {filename} to \n\t\t{destination_path}")
        else:
            # print(f"\t - [{get_file_size(filename)}] copying {filename} to \n\t\t{destination_path}")

            sys.stdout.write('\rcopying ' + filename + '                        ')
            sys.stdout.flush()


    if not os.path.isdir(os.path.dirname(destination_path)):
        try:
            os.makedirs(os.path.dirname(destination_path))
        except:
            pass

    copyfile(filename, destination_path)
    if delete_source:
        try:
            os.remove(filename)
        except:
            error('coule not remove {fl}, make sure it is not in use'.format(fl=filename))

def copy_directory_tree(src, dst):
    if not os.path.exists(dst):
        os.makedirs(dst)
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            shutil.copytree(s, d, dirs_exist_ok=True)
        else:
            shutil.copy2(s, d)

def remove_header_duplicates(file_path, line_index = 0, separator = None):
    header = read_from(file_path)[line_index]
    if not separator is None:
        header = header.split(separator)
    else:
        header = header.split()
    # Rename duplicate columns
    column_counts = defaultdict(int)
    modified_header = []
    for col_name in header:
        column_counts[col_name] += 1
        if column_counts[col_name] > 1:
            modified_header.append(f"{col_name}_{column_counts[col_name]}")
        else:
            modified_header.append(col_name)
    return modified_header


def get_swat_timeseries(file_path, header_index = 1, col_name = "flo_out", object_number = None, object_name = None, separator = None, skip_rows = 3):

    if not exists(file_path):
        print('! swat+ result file does not exist')
        return None
    
    modified_header = remove_header_duplicates(file_path, header_index, separator) + ['extra1', 'extra2']


    try:
        df              = pandas.read_csv(file_path, delim_whitespace = True, skiprows = skip_rows, names = modified_header, index_col=False)
    except:
        sys.stdout.write(f'\r! could not read {file_path} using pandas, check the number of columns              \n')
        sys.stdout.flush()
        return None

    df = df.drop(columns='extra1')
    df = df.drop(columns='extra2')

    # print(file_path)
    # print(df)
    df['date']  = pandas.to_datetime(dict(year=df.yr, month=df.mon, day=df.day))

    if not object_number is None:
        df          = df[df['unit']== object_number]
    
    if not object_name is None:
        df          = df[df['name']== object_name]
    
    if (not col_name is None) and (not col_name == "*"):
        df      = df[['date', col_name]]

    return df


def clip_raster(dstDS, srcDS, cutline):
    if not os.path.isdir(os.path.dirname(dstDS)): os.makedirs(os.path.dirname(dstDS))
    options = gdal.WarpOptions(cutlineLayer = f'{cutline}', multithread = True, cropToCutline = True)
    gdal.Warp(dstDS, srcDS, options = options)
    return True

this_dir = os.getcwd()

def copy_projection(input_proj_tif, out_tif):
    dataset = gdal.Open(input_proj_tif)
    projection = dataset.GetProjection()
    geotransform = dataset.GetGeoTransform()
    if projection is None and geotransform is None:
        print('No projection or geotransform found on file' + input)
        sys.exit(1)
    dataset2 = gdal.Open(out_tif, gdal.GA_Update)
    if dataset2 is None:
        print('Unable to open', out_tif, 'for writing')
        sys.exit(1)
    if geotransform is not None:
        dataset2.SetGeoTransform(geotransform)
    if projection is not None:
        dataset2.SetProjection(projection)


def set_tif_nodata(in_tif, out_tif, no_data=-999, stealth = True):
    create_path(os.path.dirname(out_tif))
    create_path("tmp")
    hide_folder("tmp")
    if stealth:
        os.system("gdal_translate -of GTiff -a_nodata {2} {0} {1} > tmp/tmp.txt".format(in_tif, out_tif, no_data))
    else:
        os.system("gdal_translate -of GTiff -a_nodata {2} {0} {1}".format(in_tif, out_tif, no_data))


def clip_features(mask, input_feature, output_feature, keep_only_types = None, v = False) -> geopandas.GeoDataFrame:
    '''
    keep_only_types = ['MultiPolygon', 'Polygon', 'Point', etc]
    
    '''
    mask_gdf = geopandas.read_file(mask)
    input_gdf = geopandas.read_file(input_feature)
    create_path(output_feature)
    out_gdf = input_gdf.clip(mask_gdf.to_crs(input_gdf.crs))

    if not keep_only_types is None:
        out_gdf = out_gdf[out_gdf.geometry.apply(lambda x : x.type in keep_only_types)]

    out_gdf.to_file(output_feature)

    if v:
        print("\t  - clipped feature to " + output_feature)
    return out_gdf


def assign_default_projection(in_raster, out_raster, min_lon=-180, max_lon=180, min_lat=-90, max_lat=90):
    os.system("gdal_translate -a_srs WGS84 -a_ullr {min_lon} {max_lat} {max_lon} {min_lat} {in_raster} {out_raster}".format(
        min_lon=min_lon,
        max_lon=max_lon,
        min_lat=min_lat,
        max_lat=max_lat,
        in_raster=in_raster,
        out_raster=out_raster,
    ))
